from docx import Document
from pathlib import Path
import os

MANUSCRIPT_DIR = Path("outputs/manuscript")


def export_manuscript(output_path="outputs/manuscript.docx"):
    doc = Document()

    doc.add_heading("LFX Research Copilot Manuscript", 0)

    files = [
        ("Introduction", "introduction.md"),
        ("Literature Review", "literature_review.md"),
        ("Methods", "methods_draft.md"),
        ("Discussion", "discussion_draft.md"),
        ("References", "references.md"),
    ]

    for title, filename in files:

        filepath = MANUSCRIPT_DIR / filename

        if filepath.exists():

            doc.add_heading(title, level=1)

            text = filepath.read_text(
                encoding="utf-8"
            )

            doc.add_paragraph(text)

    os.makedirs("outputs", exist_ok=True)

    doc.save(output_path)

    print(f"DOCX saved: {output_path}")


if __name__ == "__main__":
    export_manuscript()
